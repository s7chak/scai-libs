import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.tsa.seasonal import seasonal_decompose
from statsmodels.tsa.stattools import adfuller
import ops.MLConfig as MLConfig

class Explorer:
    def __init__(self, data_prepper):
        self.data_prepper = data_prepper
        self.raw_df = data_prepper.raw_df
        self.transformed_df = data_prepper.transformed_df
        self.var_metadata = data_prepper.var_metadata
        self.report = Document()

    def generate_eda_report(self, output_filename):
        self.data_summary(self.raw_df, "Data Summary")
        self.data_summary(self.transformed_df, "Transformed Data Summary")
        self.correlation(self.raw_df, "Correlation")
        self.correlation(self.transformed_df, "Transformed Correlation")
        self.vif(self.raw_df, "VIF")
        self.vif(self.transformed_df, "Transformed VIF")

        if self.data_prepper.modeling_type == MLConfig.tsa_modeltype:
            for variable, meta in self.var_metadata.items():
                if meta['ctype'] == 'continuous':
                    self.seasonality(variable)
                    self.stationarity(variable)

        for variable, meta in self.var_metadata.items():
            if meta['ctype'] == 'continuous':
                self.add_section(variable)
                self.missing_data(variable)
                self.data_plot(variable)

        self.save_report(output_filename)

    def add_section(self, variable):
        self.report.add_heading(f'Analysis for {variable}', level=1)

    def missing_data(self, variable):
        total_missing = self.raw_df[variable].null_count().sum()
        missing_percentage = (total_missing / len(self.raw_df)) * 100
        self.report.add_heading('Missing Data', level=2)
        self.report.add_paragraph(f'Total Missing Rows: {total_missing}')
        self.report.add_paragraph(f'Missing Percentage: {missing_percentage:.2f}%')

        plt.figure(figsize=(6, 4))
        sns.heatmap(self.raw_df[[variable]].is_null().to_pandas(), cbar=False)
        plt.title(f'Missing Data for {variable}')
        plt.savefig(f'missing_{variable}.png')
        self.report.add_picture(f'missing_{variable}.png')

    def data_plot(self, variable):
        self.report.add_heading('Data Plot', level=2)
        if 'Date' in self.raw_df.columns:
            plt.figure(figsize=(10, 6))
            plt.plot(self.raw_df['Date'].to_pandas(), self.raw_df[variable].to_pandas())
            plt.title(f'{variable} over Time')
            plt.savefig(f'{variable}_plot.png')
            self.report.add_picture(f'{variable}_plot.png')
        else:
            plt.figure(figsize=(10, 6))
            self.raw_df[variable].to_pandas().plot()
            plt.title(f'{variable} Plot')
            plt.savefig(f'{variable}_plot.png')
            self.report.add_picture(f'{variable}_plot.png')

    def data_summary(self, df, title):
        self.report.add_heading(title, level=2)
        summary = df.describe().to_pandas().to_dict()
        for stat, values in summary.items():
            self.report.add_paragraph(f'{stat}: {values}')

    def correlation(self, df, title):
        self.report.add_heading(f'{title}', level=2)
        corr_matrix = df.to_pandas().corr()
        plt.figure(figsize=(8, 6))
        sns.heatmap(corr_matrix, annot=True, cmap='coolwarm')
        plt.title(f'{title} Matrix')
        plt.savefig(f'{title}_correlation_matrix.png')
        self.report.add_picture(f'{title}_correlation_matrix.png')

    def vif(self, df, title):
        self.report.add_heading(f'{title} (Variance Inflation Factor)', level=2)
        X = df[self.var_metadata.keys()].to_pandas()
        vif_data = pd.DataFrame()
        vif_data["VIF"] = [variance_inflation_factor(X.values, i) for i in range(X.shape[1])]
        vif_data["Variable"] = X.columns
        self.report.add_paragraph(str(vif_data))

    def seasonality(self, variable):
        self.report.add_heading('Seasonality', level=2)
        decomposition = seasonal_decompose(self.raw_df[variable].to_pandas(), model='additive', period=12)
        decomposition.plot()
        plt.savefig(f'seasonality_{variable}.png')
        self.report.add_picture(f'seasonality_{variable}.png')

    def stationarity(self, variable):
        self.report.add_heading('Stationarity Test', level=2)
        result = adfuller(self.raw_df[variable].to_pandas())
        self.report.add_paragraph(f'ADF Statistic: {result[0]:.2f}')
        self.report.add_paragraph(f'p-value: {result[1]:.2f}')
        self.report.add_paragraph(f'Critical Values:')
        for key, value in result[4].items():
            self.report.add_paragraph(f'{key}: {value:.3f}')

    def save_report(self, output_filename):
        self.report.save(output_filename)